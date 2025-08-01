from docx import Document
from docx.shared import RGBColor, Pt


def hide_real_line(doc,message):
    para =doc.add_paragraph(message)
    run = para.runs[0]
    run.font.color.rgb = RGBColor(255,255,255)
    return para

def set_spacing(para):
    para_format = para.paragraph_format
    para_format.space_before = Pt(0)
    para_format.space_after = Pt(0)

fake = Document("C:\\Users\\ashis\\OneDrive\\Desktop\\Python_learning\\Chapter 6\\fakeMessage.docx")
real = Document("C:\\Users\\ashis\\OneDrive\\Desktop\\Python_learning\\Chapter 6\\realMessage.docx")
template = Document("C:\\Users\\ashis\\OneDrive\\Desktop\\Python_learning\\Chapter 6\\template.docx")

fake_list=[]
for line in fake.paragraphs:
    fake_list.append(line.text)

real_list=[]
for line in real.paragraphs:
    real_list.append(line.text)

require_doc = template
require_doc.add_heading('Morland Holmes', 0)
subtitle = require_doc.add_heading('Global Consultanting & Negotiations', 1)
subtitle.alignment = 1 
require_doc.add_heading('', 1)
require_doc.add_paragraph('December 17, 2015')
require_doc.add_paragraph('')

length_real_list = len(real_list)
count=0

for line in fake_list:
    if count<length_real_list and line=="":
        para = hide_real_line(require_doc,real_list[count])
        
        count += 1
    else:
        para = require_doc.add_paragraph(line)
    set_spacing(para)


require_doc.save("ciphertext_message_letterhead_1.docx")